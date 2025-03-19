import docx
import os

fpath=r'C:\Users\Jackb\OneDrive\Documents\Testing Doc.docx'
doc = docx.Document(fpath)
print(len(doc.paragraphs))
for paragraph in (doc.paragraphs):
    print(paragraph.text)

